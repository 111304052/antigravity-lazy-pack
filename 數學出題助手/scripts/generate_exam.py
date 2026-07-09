# -*- coding: utf-8 -*-
import os
import sys
from docx import Document
from docx.shared import Pt, Inches

def create_mock_exam(output_path):
    doc = Document()
    
    # 設置標題
    title = doc.add_paragraph()
    run = title.add_run("國中數學 一元一次方程式 模擬考卷")
    run.font.name = 'Microsoft JhengHei'
    run.font.size = Pt(16)
    run.bold = True
    
    # 考試資訊
    info = doc.add_paragraph()
    run_info = info.add_run("班級：______  座號：______  姓名：______  得分：______")
    run_info.font.name = 'Microsoft JhengHei'
    run_info.font.size = Pt(11)
    
    doc.add_paragraph("\n一、 選擇題（每題 20 分，共 100 分）")
    
    questions = [
        "1. 解方程式 2x + 5 = 15，求 x 的值？\n   (A) 3  (B) 5  (C) 10  (D) 20",
        "2. 若 3(x - 2) = 12，則 x 的值為何？\n   (A) 2  (B) 4  (C) 6  (D) 8",
        "3. 媽媽買了 3 顆蘋果 and 1 個 50 元的禮盒，總共花了 140 元。若每顆蘋果 x 元，可列出下列哪一個方程式？\n   (A) 3x - 50 = 140  (B) 3x + 50 = 140  (C) 3 + 50x = 140  (D) 3x = 140",
        "4. 解方程式 5x - 7 = 2x + 8，求 x 的值？\n   (A) 3  (B) 5  (C) 7  (D) 9",
        "5. 小明今年 x 歲，爸爸的年齡比小明的 3 倍多 2 歲。若爸爸今年 38 歲，則小明今年幾歲？\n   (A) 10 歲  (B) 11 歲  (C) 12 歲  (D) 13 歲"
    ]
    
    for q in questions:
        p = doc.add_paragraph()
        run_q = p.add_run(q)
        run_q.font.name = 'Microsoft JhengHei'
        run_q.font.size = Pt(11)
        p.paragraph_format.left_indent = Inches(0.2)
        
    doc.add_paragraph("\n二、 參考解答：\n1. (B)   2. (C)   3. (B)   4. (B)   5. (C)")
    
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")
    print(f"[SUCCESS] 考卷生成成功：{output_path}")

if __name__ == "__main__":
    out_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", "output"))
    out_file = os.path.join(out_dir, "數學模擬考卷_一元一次方程式.docx")
    create_mock_exam(out_file)
